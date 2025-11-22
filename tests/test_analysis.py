from pathlib import Path

from docx import Document

from acm.analysis import analyze_manuscript, journal_change_requests, parse_docx_sections
from acm.journal import Guideline


def _build_doc(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Intro words here.")
    doc.add_paragraph("More intro text for counting.")

    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Method details are written here.")

    doc.add_heading("Results", level=1)
    doc.add_paragraph("Results are short.")

    path = tmp_path / "sample.docx"
    doc.save(path)
    return path


def test_parse_docx_sections_counts(tmp_path: Path) -> None:
    path = _build_doc(tmp_path)

    sections = parse_docx_sections(path)

    assert [s.title for s in sections] == ["Introduction", "Methods", "Results"]
    assert [s.word_count for s in sections] == [8, 5, 3]
    assert [s.category for s in sections] == ["Introduction", "Methods", "Results"]


def test_analyze_manuscript_matches_guidelines(tmp_path: Path) -> None:
    path = _build_doc(tmp_path)
    guidelines = [
        Guideline(
            journal="Journal A",
            article_type="Research",
            word_limit="20 words",
            structure="Introduction, Methods, Results",
        ),
        Guideline(
            journal="Journal B",
            article_type="Research",
            word_limit="10 words",
            structure="Introduction and Discussion",
        ),
    ]

    result = analyze_manuscript(path, guidelines)

    assert "Journal A" in result.accepted_journals
    assert "Journal B" in result.required_changes
    journal_b_changes = " ".join(result.required_changes["Journal B"])
    assert "Reduce word count" in journal_b_changes
    assert "Discussion" in journal_b_changes


def test_journal_change_requests_include_abstract_limit(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("This abstract section is intentionally too long for testing.")
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Intro text")
    path = tmp_path / "abstract.docx"
    doc.save(path)

    guideline = Guideline(
        journal="Journal C",
        article_type="Research",
        abstract_limit="5 words",
    )

    sections = parse_docx_sections(path)
    changes = journal_change_requests(guideline, sections)

    assert any("Abstract exceeds limit" in change for change in changes)
