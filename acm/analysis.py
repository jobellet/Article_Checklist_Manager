from __future__ import annotations

"""Utilities for analysing manuscript files and matching journal guidelines."""

from dataclasses import dataclass
import re
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Set

from docx import Document

from .journal import Guideline, load_guidelines


SECTION_KEYWORDS: Dict[str, str] = {
    "introduction": "Introduction",
    "background": "Introduction",
    "method": "Methods",
    "methods": "Methods",
    "materials and methods": "Methods",
    "methodology": "Methods",
    "result": "Results",
    "results": "Results",
    "discussion": "Discussion",
    "conclusion": "Conclusion",
    "conclusions": "Conclusion",
    "abstract": "Abstract",
}


@dataclass
class SectionSummary:
    """Summarised section extracted from a manuscript file."""

    title: str
    word_count: int
    category: str


@dataclass
class AnalysisResult:
    """Result of analysing a manuscript against journal guidelines."""

    sections: List[SectionSummary]
    total_words: int
    categories: Set[str]
    accepted_journals: List[str]
    required_changes: Dict[str, List[str]]


def _is_heading(paragraph) -> bool:
    style = paragraph.style.name if paragraph.style else ""
    return style.lower().startswith("heading")


def categorize_section(title: str) -> str:
    """Return a normalised section category inferred from ``title``."""

    lowered = title.lower()
    for key, category in SECTION_KEYWORDS.items():
        if key in lowered:
            return category
    return "Other"


def parse_docx_sections(path: Path) -> List[SectionSummary]:
    """Parse a ``.docx`` file into section summaries.

    Sections are defined as the text between heading paragraphs.
    """

    document = Document(path)
    sections: List[SectionSummary] = []
    current_title: Optional[str] = None
    word_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if _is_heading(paragraph):
            if current_title is not None:
                sections.append(
                    SectionSummary(
                        title=current_title,
                        word_count=word_count,
                        category=categorize_section(current_title),
                    )
                )
            current_title = text
            word_count = 0
        else:
            word_count += len(text.split())

    if current_title is None:
        total_words = sum(len(p.text.split()) for p in document.paragraphs if p.text.strip())
        return [SectionSummary(title="Document", word_count=total_words, category="Other")]

    sections.append(
        SectionSummary(
            title=current_title,
            word_count=word_count,
            category=categorize_section(current_title),
        )
    )
    return sections


def _parse_word_limit(limit: Optional[str]) -> Optional[int]:
    if not limit:
        return None
    match = re.search(r"(\d[\d,]*)", limit)
    if not match:
        return None
    return int(match.group(1).replace(",", ""))


def _required_categories(structure: Optional[str]) -> Set[str]:
    if not structure:
        return set()
    lower = structure.lower()
    categories: Set[str] = set()
    for key, category in SECTION_KEYWORDS.items():
        if key in lower:
            categories.add(category)
    return categories


def _journal_fit(
    guideline: Guideline,
    sections: Sequence[SectionSummary],
    categories: Set[str],
    total_words: int,
) -> tuple[bool, List[str]]:
    changes: List[str] = []
    required_sections = _required_categories(guideline.structure)
    missing = sorted(required_sections - categories)
    if missing:
        changes.append(
            "Add sections covering: " + ", ".join(missing)
        )

    limit = _parse_word_limit(guideline.word_limit)
    if limit is not None and total_words > limit:
        changes.append(
            f"Reduce word count by {total_words - limit} to meet {limit}-word limit"
        )

    abstract_limit = _parse_word_limit(guideline.abstract_limit)
    if abstract_limit is not None:
        abstract_words = next(
            (section.word_count for section in sections if section.category == "Abstract"),
            0,
        )
        if abstract_words > abstract_limit:
            changes.append(
                (
                    "Abstract exceeds limit: "
                    f"{abstract_words}/{abstract_limit} words (reduce by {abstract_words - abstract_limit})"
                )
            )

    return not changes, changes


def journal_change_requests(
    guideline: Guideline, sections: Sequence[SectionSummary]
) -> List[str]:
    """Return change requests for ``guideline`` based on manuscript ``sections``."""

    categories = {section.category for section in sections if section.category != "Other"}
    total_words = sum(section.word_count for section in sections)
    _, changes = _journal_fit(guideline, sections, categories, total_words)
    return changes


def analyze_manuscript(
    path: Path, guidelines: Optional[Iterable[Guideline]] = None
) -> AnalysisResult:
    """Analyse a manuscript file and compare it with journal guidelines."""

    sections = parse_docx_sections(path)
    categories = {section.category for section in sections if section.category != "Other"}
    total_words = sum(section.word_count for section in sections)

    all_guidelines = list(guidelines) if guidelines is not None else load_guidelines()
    accepted: List[str] = []
    changes_needed: Dict[str, List[str]] = {}

    for guideline in all_guidelines:
        fits, changes = _journal_fit(guideline, sections, categories, total_words)
        if fits:
            accepted.append(guideline.journal)
        elif changes:
            changes_needed[guideline.journal] = changes

    return AnalysisResult(
        sections=sections,
        total_words=total_words,
        categories=categories,
        accepted_journals=accepted,
        required_changes=changes_needed,
    )


__all__ = [
    "SectionSummary",
    "AnalysisResult",
    "categorize_section",
    "parse_docx_sections",
    "analyze_manuscript",
    "journal_change_requests",
]
